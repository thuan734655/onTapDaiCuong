from docx import Document
import os

doc_path = os.path.join(os.path.dirname(__file__), 'chủ nghĩa.docx')
doc = Document(doc_path)

# Xuất ra file text để phân tích
with open('content_analysis.txt', 'w', encoding='utf-8') as f:
    f.write(f"Tổng số paragraphs: {len(doc.paragraphs)}\n\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"[{i}] {para.text}\n")

print("Đã xuất nội dung ra file content_analysis.txt")
